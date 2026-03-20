"""
Convert WORKFLOW.md to DOCX and PDF.
Requires: pip install python-docx markdown
PDF: Uses weasyprint if available, else suggests md-to-pdf via npx.
"""
import re
import os

def convert_to_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Installing python-docx...")
        os.system("pip install python-docx -q")
        from docx import Document
        from docx.shared import Pt, Inches

    doc = Document()
    doc.add_heading("Miisky Project — Workflow Overview", 0)

    with open("WORKFLOW.md", "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip title line (already added as heading)
        if stripped == "# Miisky Project — Workflow Overview":
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.add_run("\n".join(code_lines)).font.name = "Consolas"
                p.paragraph_format.left_indent = Inches(0.25)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if "|" in line and stripped.startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                # Parse table, skip separator row (|---|---|)
                parsed = []
                for r in table_lines:
                    if not r.strip():
                        continue
                    cells = [c.strip() for c in r.split("|")[1:-1]]
                    if cells and not re.match(r"^[\s\-:]+$", "".join(cells)):
                        parsed.append(cells)
                if parsed:
                    col_count = max(len(r) for r in parsed)
                    t = doc.add_table(rows=len(parsed), cols=col_count)
                    t.style = "Table Grid"
                    for ri, row in enumerate(parsed):
                        for ci, cell in enumerate(row[:col_count]):
                            t.rows[ri].cells[ci].text = cell
                table_lines = []
            in_table = False

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Horizontal rule
        elif stripped == "---":
            doc.add_paragraph("—" * 20)
        # List
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(stripped, style="List Number")
        # Empty
        elif not stripped:
            pass
        # Normal paragraph
        else:
            # Bold
            text = stripped
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2] + " ")
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    # Flush any remaining
    if in_code_block and code_lines:
        p = doc.add_paragraph()
        p.add_run("\n".join(code_lines)).font.name = "Consolas"
    if in_table and table_lines:
        parsed = []
        for r in table_lines:
            if not r.strip():
                continue
            cells = [c.strip() for c in r.split("|")[1:-1]]
            if cells and not re.match(r"^[\s\-:]+$", "".join(cells)):
                parsed.append(cells)
        if parsed:
            col_count = max(len(r) for r in parsed)
            t = doc.add_table(rows=len(parsed), cols=col_count)
            t.style = "Table Grid"
            for ri, row in enumerate(parsed):
                for ci, cell in enumerate(row[:col_count]):
                    t.rows[ri].cells[ci].text = cell

    out_path = os.path.join(os.path.dirname(__file__), "WORKFLOW.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")
    return out_path


def convert_to_pdf():
    """Convert DOCX to PDF using docx2pdf (Windows: uses Word) or fallback to npx md-to-pdf."""
    docx_path = os.path.join(os.path.dirname(__file__), "WORKFLOW.docx")
    pdf_path = os.path.join(os.path.dirname(__file__), "WORKFLOW.pdf")
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"Created: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"docx2pdf failed ({e}). Try: open WORKFLOW.docx in Word → Save As PDF")
        return None


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    print("Converting to DOCX...")
    convert_to_docx()
    print("Converting to PDF...")
    convert_to_pdf()
